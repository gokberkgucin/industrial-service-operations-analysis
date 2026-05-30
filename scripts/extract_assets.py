from __future__ import annotations

import csv
import shutil
import traceback
import zipfile
from datetime import datetime
from pathlib import Path

try:
    import fitz  # PyMuPDF
except Exception as exc:  # pragma: no cover
    fitz = None
    FITZ_IMPORT_ERROR = exc
else:
    FITZ_IMPORT_ERROR = None

try:
    from docx import Document
except Exception as exc:  # pragma: no cover
    Document = None
    DOCX_IMPORT_ERROR = exc
else:
    DOCX_IMPORT_ERROR = None


REPO_ROOT = Path(__file__).resolve().parents[1]

DOCX_PATH = Path(
    r"D:\TOPLANDIK\IK\Teknik_disi_arka_plan_ve_ilk_donem_proje_belgeleri\EK-10_Siemens_Staj_Makale_Analizleri.docx"
)
ARTICLE_1_PATH = Path(
    r"D:\TOPLANDIK\IK\Teknik_disi_arka_plan_ve_ilk_donem_proje_belgeleri\Chase - A history of research in service operations  What s the big idea.pdf"
)
ARTICLE_2_PATH = Path(
    r"D:\TOPLANDIK\IK\Teknik_disi_arka_plan_ve_ilk_donem_proje_belgeleri\Cross-training policies in field services.pdf"
)

EXTRACTED_ROOT = REPO_ROOT / "_work" / "extracted"
LOG_DIR = REPO_ROOT / "_work" / "logs"
BACKUP_ROOT = REPO_ROOT / "_work" / "backup" / f"extraction_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
SUMMARY_PATH = REPO_ROOT / "metadata" / "extraction-summary.md"
ERROR_LOG_PATH = LOG_DIR / "extraction-errors.md"

ERRORS: list[str] = []
BACKED_UP_PATHS: list[str] = []


def log_error(context: str, exc: BaseException) -> None:
    ERRORS.append(f"## {context}\n\n```text\n{traceback.format_exc().strip()}\n```\n")


def relative_display(path: Path) -> str:
    try:
        return str(path.relative_to(REPO_ROOT)).replace("\\", "/")
    except ValueError:
        return str(path)


def backup_if_exists(path: Path) -> None:
    if not path.exists():
        return
    try:
        relative = path.relative_to(REPO_ROOT)
    except ValueError:
        relative = Path(path.name)
    destination = BACKUP_ROOT / relative
    destination.parent.mkdir(parents=True, exist_ok=True)
    if path.is_dir():
        shutil.copytree(path, destination, dirs_exist_ok=True)
    else:
        shutil.copy2(path, destination)
    BACKED_UP_PATHS.append(relative_display(path))


def reset_directory(path: Path) -> None:
    backup_if_exists(path)
    if path.exists():
        shutil.rmtree(path)
    path.mkdir(parents=True, exist_ok=True)


def write_text_with_backup(path: Path, content: str) -> None:
    backup_if_exists(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def write_bytes(path: Path, content: bytes) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(content)


def ensure_sources_exist() -> None:
    missing = [path for path in (DOCX_PATH, ARTICLE_1_PATH, ARTICLE_2_PATH) if not path.exists()]
    if missing:
        formatted = "\n".join(f"- {path}" for path in missing)
        raise FileNotFoundError(f"Kaynak dosya bulunamadı:\n{formatted}")


def extract_docx_text() -> int:
    output_path = EXTRACTED_ROOT / "report-raw.txt"
    lines: list[str] = []

    if Document is None:
        raise RuntimeError(f"python-docx yüklenemedi: {DOCX_IMPORT_ERROR}")

    document = Document(DOCX_PATH)
    for paragraph in document.paragraphs:
        lines.append(paragraph.text)

    if document.tables:
        lines.append("")
        lines.append("[TABLOLAR]")
        for table_index, table in enumerate(document.tables, start=1):
            lines.append(f"[TABLO {table_index}]")
            for row in table.rows:
                cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                lines.append("\t".join(cells))

    write_text_with_backup(output_path, "\n".join(lines).strip() + "\n")
    return 1


def extract_docx_images() -> int:
    image_dir = EXTRACTED_ROOT / "docx-images"
    reset_directory(image_dir)
    count = 0
    with zipfile.ZipFile(DOCX_PATH) as archive:
        for member in archive.infolist():
            if not member.filename.startswith("word/media/") or member.is_dir():
                continue
            output_path = image_dir / Path(member.filename).name
            write_bytes(output_path, archive.read(member))
            count += 1
    return count


def extract_pdf(pdf_path: Path, output_dir: Path) -> dict[str, int]:
    if fitz is None:
        raise RuntimeError(f"PyMuPDF yüklenemedi: {FITZ_IMPORT_ERROR}")

    reset_directory(output_dir)
    page_images_dir = output_dir / "page-images"
    embedded_images_dir = output_dir / "embedded-images"
    page_images_dir.mkdir(parents=True, exist_ok=True)
    embedded_images_dir.mkdir(parents=True, exist_ok=True)

    counts = {
        "page_text_files": 0,
        "rendered_page_images": 0,
        "embedded_images": 0,
    }
    manifest_rows: list[dict[str, str | int]] = []
    full_text_parts: list[str] = []

    with fitz.open(pdf_path) as document:
        for page_number, page in enumerate(document, start=1):
            page_label = f"page-{page_number:03d}"

            text = page.get_text("text") or ""
            write_text_with_backup(output_dir / f"{page_label}.txt", text)
            full_text_parts.append(f"\n\n--- {page_label} ---\n\n{text}")
            counts["page_text_files"] += 1

            try:
                pixmap = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
                pixmap.save(str(page_images_dir / f"{page_label}.png"))
                counts["rendered_page_images"] += 1
            except Exception as exc:
                log_error(f"{pdf_path.name} {page_label} sayfa görüntüsü", exc)

            try:
                for image_index, image in enumerate(page.get_images(full=True), start=1):
                    xref = image[0]
                    extracted = document.extract_image(xref)
                    image_bytes = extracted.get("image")
                    extension = extracted.get("ext", "bin")
                    if not image_bytes:
                        continue
                    image_name = f"{page_label}-image-{image_index:02d}.{extension}"
                    image_path = embedded_images_dir / image_name
                    write_bytes(image_path, image_bytes)
                    counts["embedded_images"] += 1
                    manifest_rows.append(
                        {
                            "page": page_number,
                            "image_index": image_index,
                            "xref": xref,
                            "extension": extension,
                            "output_path": relative_display(image_path),
                        }
                    )
            except Exception as exc:
                log_error(f"{pdf_path.name} {page_label} gömülü görsel çıkarımı", exc)

    write_text_with_backup(output_dir / "full-text.txt", "".join(full_text_parts).strip() + "\n")

    manifest_path = output_dir / "figure-info.csv"
    with manifest_path.open("w", encoding="utf-8", newline="") as csv_file:
        writer = csv.DictWriter(csv_file, fieldnames=["page", "image_index", "xref", "extension", "output_path"])
        writer.writeheader()
        writer.writerows(manifest_rows)

    return counts


def write_error_log() -> None:
    LOG_DIR.mkdir(parents=True, exist_ok=True)
    if ERRORS:
        content = "# Extraction Errors\n\n" + "\n".join(ERRORS)
    else:
        content = "# Extraction Errors\n\nHata kaydı yok.\n"
    write_text_with_backup(ERROR_LOG_PATH, content)


def write_summary(counts: dict[str, int | str]) -> None:
    output_dirs = [
        "_work/extracted/",
        "_work/extracted/docx-images/",
        "_work/extracted/article-1/",
        "_work/extracted/article-1/page-images/",
        "_work/extracted/article-1/embedded-images/",
        "_work/extracted/article-2/",
        "_work/extracted/article-2/page-images/",
        "_work/extracted/article-2/embedded-images/",
    ]
    backup_text = "\n".join(f"- {path}" for path in BACKED_UP_PATHS) if BACKED_UP_PATHS else "- Yedeklenen eski çıktı yok."
    error_text = f"{len(ERRORS)} hata kaydı var; ayrıntı için `_work/logs/extraction-errors.md`." if ERRORS else "Hata kaydı yok."
    summary = f"""# Extraction Summary

Tarih: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

Bu dosya, yerel DOCX raporundan ve iki PDF makaleden referans amaçlı ham çıkarım üretildiğini özetler. Ham kaynak dosyalar repoya kopyalanmamıştır.

## Çıktı Sayıları

| Çıktı | Sayı |
|---|---:|
| DOCX ham metin dosyası | {counts['docx_text_files']} |
| DOCX görsel dosyası | {counts['docx_images']} |
| Makale 1 sayfa metin dosyası | {counts['article_1_page_text_files']} |
| Makale 1 sayfa görüntüsü | {counts['article_1_rendered_page_images']} |
| Makale 1 gömülü görsel | {counts['article_1_embedded_images']} |
| Makale 2 sayfa metin dosyası | {counts['article_2_page_text_files']} |
| Makale 2 sayfa görüntüsü | {counts['article_2_rendered_page_images']} |
| Makale 2 gömülü görsel | {counts['article_2_embedded_images']} |

## Oluşan Klasörler

{chr(10).join(f"- `{directory}`" for directory in output_dirs)}

## Yedekleme

Var olan çıktıların üzerine yazılmadan önce yedekleme mantığı çalıştırıldı.

{backup_text}

## Hata Durumu

{error_text}

## Not

`_work/` altındaki çıkarımlar ham çalışma malzemesidir. Kamuya açık içerik olarak doğrudan kullanılmamalı; sonraki aşamada seçilmeli, özetlenmeli ve telif/gizlilik açısından temizlenmelidir.
"""
    write_text_with_backup(SUMMARY_PATH, summary)


def main() -> int:
    ensure_sources_exist()
    (EXTRACTED_ROOT).mkdir(parents=True, exist_ok=True)
    (REPO_ROOT / "metadata").mkdir(parents=True, exist_ok=True)

    counts: dict[str, int | str] = {
        "docx_text_files": 0,
        "docx_images": 0,
        "article_1_page_text_files": 0,
        "article_1_rendered_page_images": 0,
        "article_1_embedded_images": 0,
        "article_2_page_text_files": 0,
        "article_2_rendered_page_images": 0,
        "article_2_embedded_images": 0,
    }

    try:
        counts["docx_text_files"] = extract_docx_text()
    except Exception as exc:
        log_error("DOCX metin çıkarımı", exc)

    try:
        counts["docx_images"] = extract_docx_images()
    except Exception as exc:
        log_error("DOCX görsel çıkarımı", exc)

    try:
        article_1_counts = extract_pdf(ARTICLE_1_PATH, EXTRACTED_ROOT / "article-1")
        counts["article_1_page_text_files"] = article_1_counts["page_text_files"]
        counts["article_1_rendered_page_images"] = article_1_counts["rendered_page_images"]
        counts["article_1_embedded_images"] = article_1_counts["embedded_images"]
    except Exception as exc:
        log_error("Makale 1 PDF çıkarımı", exc)

    try:
        article_2_counts = extract_pdf(ARTICLE_2_PATH, EXTRACTED_ROOT / "article-2")
        counts["article_2_page_text_files"] = article_2_counts["page_text_files"]
        counts["article_2_rendered_page_images"] = article_2_counts["rendered_page_images"]
        counts["article_2_embedded_images"] = article_2_counts["embedded_images"]
    except Exception as exc:
        log_error("Makale 2 PDF çıkarımı", exc)

    write_error_log()
    write_summary(counts)

    print("Extraction tamamlandı.")
    print(f"Özet: {SUMMARY_PATH}")
    print(f"Hata logu: {ERROR_LOG_PATH}")
    return 0 if not ERRORS else 1


if __name__ == "__main__":
    raise SystemExit(main())
