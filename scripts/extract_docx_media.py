#!/usr/bin/env python
"""Extract raw DOCX media and transcript inputs without modifying the source DOCX."""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import shutil
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from PIL import Image, UnidentifiedImageError


def repo_root() -> Path:
    return Path(__file__).resolve().parents[1]


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def guess_kind(extension: str, width: int | None, height: int | None) -> str:
    ext = extension.lower()
    if ext == ".emf":
        return "vector-or-office-drawing"
    if ext == ".svg":
        return "vector"
    if width is None or height is None:
        return "unknown"
    if width < 120 or height < 120:
        return "small-decorative-or-icon"
    if width > 900 or height > 700:
        return "large-figure-or-screenshot"
    return "embedded-image"


def iter_block_text(document: Document) -> list[dict[str, object]]:
    blocks: list[dict[str, object]] = []
    block_index = 1

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(
                {
                    "index": block_index,
                    "type": "paragraph",
                    "style": paragraph.style.name if paragraph.style else "",
                    "text": text,
                }
            )
            block_index += 1

    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        blocks.append(
            {
                "index": block_index,
                "type": "table",
                "style": "",
                "table_index": table_index,
                "rows": rows,
            }
        )
        block_index += 1

    return blocks


def main() -> int:
    root = repo_root()
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--docx",
        default=str(root / "EK-10_Siemens_Staj_Makale_Analizleri.docx"),
        help="Source DOCX path.",
    )
    parser.add_argument(
        "--out",
        default=str(root / "_work" / "extracted" / "docx"),
        help="Extraction output directory.",
    )
    args = parser.parse_args()

    docx_path = Path(args.docx).resolve()
    out_dir = Path(args.out).resolve()
    media_dir = out_dir / "media_raw"
    log_dir = root / "_work" / "logs"

    if not docx_path.exists():
        raise FileNotFoundError(f"Source DOCX not found: {docx_path}")

    media_dir.mkdir(parents=True, exist_ok=True)
    log_dir.mkdir(parents=True, exist_ok=True)

    inventory_rows = []
    errors = []

    with zipfile.ZipFile(docx_path) as archive:
        media_names = sorted(name for name in archive.namelist() if name.startswith("word/media/") and not name.endswith("/"))
        for media_name in media_names:
            source_name = Path(media_name).name
            output_path = media_dir / source_name
            with archive.open(media_name) as source, output_path.open("wb") as target:
                shutil.copyfileobj(source, target)

            width = None
            height = None
            error = ""
            try:
                with Image.open(output_path) as image:
                    width, height = image.size
            except (UnidentifiedImageError, OSError) as exc:
                error = str(exc)
                errors.append({"file": source_name, "error": error})

            extension = output_path.suffix.lower()
            inventory_rows.append(
                {
                    "original_name": source_name,
                    "extension": extension,
                    "bytes": output_path.stat().st_size,
                    "width": width or "",
                    "height": height or "",
                    "sha256": sha256(output_path),
                    "kind": guess_kind(extension, width, height),
                    "output_path": str(output_path.relative_to(root)).replace("\\", "/"),
                    "read_error": error,
                }
            )

    document = Document(docx_path)
    blocks = iter_block_text(document)
    transcript_lines = []
    for block in blocks:
        if block["type"] == "paragraph":
            transcript_lines.append(str(block["text"]))
        elif block["type"] == "table":
            transcript_lines.append(f"[TABLE {block['table_index']}]")
            for row in block["rows"]:
                transcript_lines.append(" | ".join(row))

    (out_dir / "transcript_raw.txt").write_text("\n\n".join(transcript_lines), encoding="utf-8")
    (out_dir / "blocks_raw.json").write_text(json.dumps(blocks, ensure_ascii=False, indent=2), encoding="utf-8")

    inventory_path = out_dir / "docx_media_inventory.csv"
    with inventory_path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.DictWriter(
            handle,
            fieldnames=[
                "original_name",
                "extension",
                "bytes",
                "width",
                "height",
                "sha256",
                "kind",
                "output_path",
                "read_error",
            ],
        )
        writer.writeheader()
        writer.writerows(inventory_rows)

    result = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "source_docx": str(docx_path),
        "media_count": len(inventory_rows),
        "text_block_count": len(blocks),
        "inventory": str(inventory_path),
        "transcript_raw": str(out_dir / "transcript_raw.txt"),
        "blocks_raw": str(out_dir / "blocks_raw.json"),
        "errors": errors,
    }
    (log_dir / "extract_docx_media.json").write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
